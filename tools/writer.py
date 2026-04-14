"""
Document Creation Engine — create, format, export documents.

Creates professional-quality documents in multiple formats:
  - Markdown (.md) — source of truth, clean and versionable
  - PDF (.pdf) — via weasyprint, print-ready
  - DOCX (.docx) — via python-docx, editable in Word
  - HTML (.html) — clean styled HTML
  - Plain text (.txt) — simple and clean

Templates:
  - Company Research Report
  - Market Analysis Report
  - Meeting Preparation Document
  - Technical Documentation (README, API docs)
  - Proposal / Brief / Memo
  - Custom Document

Auto-organization:
  ~/Clawspan_Docs/
  ├── reports/
  ├── meetings/
  ├── technical/
  ├── proposals/
  └── briefs/
"""

from __future__ import annotations

import csv
import io
import json
import os
import re
import subprocess
from datetime import datetime
from typing import Any


# ── Content Cleaning ─────────────────────────────────────────────────────────

def _clean_content_for_doc(text: str) -> str:
    """Clean raw research content before inserting into a professional document.

    Removes:
    - Markdown sub-headers from scraped content (### Source Title)
    - Remaining UI/navigation artifacts
    - Table formatting garbage
    - Duplicate paragraphs
    - Very short fragments
    """
    if not text:
        return ""

    lines = text.split("\n")
    cleaned = []
    seen = set()

    noise_patterns = [
        "hero section", "skip to content", "javascript is disabled",
        "we have received your inquiry", "copyright", "all rights reserved",
        "powered by", "chrome web store", "firefox extension",
        "tracxn", "prospeo", "logo for", "logo of", "footer-bottom",
        "try our premium", "illustration", "icon", "flag of",
        "view details", "view all", "read more", "learn more",
        "srsltid", "utm_", "pdf illustration", "chrome_extension",
        "navbar", "mobile menu", "back to top", "scroll to top",
        "anniversary-logo", "footer-bottom-logos",
        "overviewemails", "formataboutemployees", "tech stacktrending",
        "view company", "key contacts at", "view all employees",
    ]

    for line in lines:
        stripped = line.strip()

        # Skip sub-headers from scraped content
        if re.match(r'^#{3,}\s+', stripped):
            continue

        # Skip empty
        if not stripped:
            continue

        # Skip lines that are primarily source attribution (contain URL-like patterns or source brands)
        source_brands = ["tracxn.com", "prospeo.io", "crunchbase.com", "linkedin.com/company"]
        if any(b in stripped.lower() for b in source_brands):
            continue
        # Also skip lines that are just source titles
        if re.match(r'^#{1,3}\s+.*\|.*\b(tracxn|prospeo|crunchbase|linkedin)\b', stripped, re.IGNORECASE):
            continue

        # Skip noise
        lower = stripped.lower()
        if any(p in lower for p in noise_patterns):
            continue

        # Skip URLs
        if stripped.startswith("http") or stripped.startswith("www."):
            continue

        # Skip table formatting
        if stripped in ("---", "|", "|---|", "| |"):
            continue

        # Skip very short lines (< 15 chars) unless they look like data
        if len(stripped) < 15 and not any(c.isdigit() for c in stripped):
            continue

        # Deduplicate
        key = stripped.lower().strip(".,!? ")
        if key and len(key) > 20:
            if key in seen:
                continue
            seen.add(key)

        cleaned.append(stripped)

    result = "\n".join(cleaned)
    result = re.sub(r'\n{3,}', '\n\n', result).strip()
    return result

# ── Document Storage ─────────────────────────────────────────────────────────

DOCS_DIR = os.path.expanduser("~/Clawspan_Docs")

DOC_SUBFOLDERS = {
    "report": "reports",
    "meeting": "meetings",
    "technical": "technical",
    "proposal": "proposals",
    "brief": "briefs",
    "default": "documents",
}


def _get_subfolder(doc_type: str) -> str:
    """Get subfolder for document type."""
    for key, folder in DOC_SUBFOLDERS.items():
        if key in doc_type.lower():
            return folder
    return DOC_SUBFOLDERS["default"]


def _generate_filename(title: str, doc_type: str = "document") -> str:
    """Generate a clean, date-stamped filename."""
    date_str = datetime.now().strftime("%Y-%m-%d")
    # Clean title: lowercase, replace spaces/special chars with underscores
    clean = re.sub(r'[^a-zA-Z0-9\s_-]', '', title)
    clean = re.sub(r'[\s]+', '_', clean).strip('_')
    clean = clean[:60]  # Truncate long titles
    if doc_type == "report":
        return f"{date_str}_{clean}_Report.md"
    elif doc_type == "meeting":
        return f"{date_str}_{clean}_Meeting_Prep.md"
    elif doc_type == "technical":
        return f"{clean}.md"
    return f"{date_str}_{clean}.md"


def _ensure_dirs() -> None:
    """Ensure document directories exist."""
    os.makedirs(DOCS_DIR, exist_ok=True)
    for folder in DOC_SUBFOLDERS.values():
        os.makedirs(os.path.join(DOCS_DIR, folder), exist_ok=True)


# ── Document Templates ───────────────────────────────────────────────────────

def _section_emoji(key: str) -> str:
    """Return a decorative emoji for each section type."""
    return {
        "overview": "🏢",
        "leadership": "👥",
        "financials": "💰",
        "products": "🛠️",
        "news": "📰",
        "competitors": "⚔️",
        "trends": "📈",
        "analyst_sentiment": "🔍",
    }.get(key, "📋")


def _clean_sources(sources: list[dict]) -> list[dict]:
    """Strip tracking params and platform noise from source titles/URLs."""
    cleaned = []
    for s in sources:
        title = re.sub(
            r'\s*[\-|]\s*\b(Tracxn|Prospeo|Crunchbase|LinkedIn|CB Insights|Glassdoor|IPo Platform|Enablers|Inc42)\b.*$',
            '', s.get("title", "Source"), flags=re.IGNORECASE,
        ).strip() or "Source"
        url = re.sub(r'[?&](srsltid|utm_[^&]+|ref|fbclid|gclid)=[^&]*', '', s.get("url", ""))
        cleaned.append({"title": title, "url": url})
    return cleaned


def _render_section_block(key: str, display_name: str, section: dict) -> list[str]:
    """Render a single research section as clean, structured markdown."""
    emoji = _section_emoji(key)
    lines = []

    lines.append(f"## {emoji} {display_name}")
    lines.append("")

    content = _clean_content_for_doc(section.get("content", ""))
    if content:
        # Demote any headings so they nest under ## section header
        content = re.sub(
            r'^(#{1,4})\s',
            lambda m: '#' * min(len(m.group(1)) + 2, 6) + ' ',
            content, flags=re.MULTILINE,
        )
        lines.append(content)
    else:
        lines.append("*No data available for this section.*")

    lines.append("")

    sources = _clean_sources(section.get("sources", []))
    if sources:
        lines.append("<details>")
        lines.append(f"<summary>📎 Sources ({len(sources)})</summary>")
        lines.append("")
        for i, s in enumerate(sources, 1):
            if s["url"]:
                lines.append(f"{i}. [{s['title']}]({s['url']})")
            else:
                lines.append(f"{i}. {s['title']}")
        lines.append("")
        lines.append("</details>")
        lines.append("")

    lines.append("---")
    lines.append("")
    return lines


def _extract_key_highlights(content: str, max_points: int = 5) -> list[str]:
    """Extract key bullet-point highlights from raw section content.

    Pulls sentences that contain numbers, names, or strong signal words
    — the kind of facts a reader wants at a glance.
    """
    if not content:
        return []

    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', content.replace('\n', ' '))

    signal_words = [
        r'\$[\d,]+', r'\d+%', r'\d+\s*(million|billion|thousand|M|B|K)',
        r'founded', r'launched', r'raised', r'acquired', r'announced',
        r'CEO|CTO|CFO|founder', r'revenue', r'valuation', r'funding',
        r'headquartered', r'employees', r'users', r'clients', r'platform',
        r'series [A-Z]', r'IPO', r'partnership', r'award', r'patent',
    ]
    signal_re = re.compile('|'.join(signal_words), re.IGNORECASE)

    scored = []
    for s in sentences:
        s = s.strip()
        if len(s) < 40 or len(s) > 300:
            continue
        # Skip lines that are just headers or noise
        if s.startswith('#') or s.startswith('|') or s.startswith('>'):
            continue
        score = len(signal_re.findall(s))
        if score > 0:
            scored.append((score, s))

    scored.sort(key=lambda x: -x[0])
    seen = set()
    result = []
    for _, s in scored:
        key = s[:60].lower()
        if key not in seen:
            seen.add(key)
            result.append(s)
        if len(result) >= max_points:
            break

    return result


def _build_smart_exec_summary(entity: str, sections: dict,
                               report_type: str = "company") -> str:
    """Build a structured, readable executive summary from research sections.

    Produces 4 blocks:
      1. What it is / overview paragraph
      2. Key numbers at a glance (table)
      3. Notable highlights (bullets)
      4. Analyst view / recent developments
    """
    lines = []

    # ── 1. What it is ───────────────────────────────────────────────────────
    overview_content = ""
    for key in ("overview",):
        if key in sections and sections[key].get("content"):
            raw = sections[key]["content"]
            # Take the first clean paragraph of substance
            for para in raw.split("\n"):
                para = para.strip()
                if len(para) > 80 and not para.startswith('#') and not para.startswith('|'):
                    overview_content = para[:600]
                    break

    if overview_content:
        lines.append(overview_content)
        lines.append("")

    # ── 2. Key Numbers ──────────────────────────────────────────────────────
    metric_patterns = [
        (r'\$[\d,.]+\s*(billion|million|B|M)\b', "Valuation / Revenue"),
        (r'Series [A-Z]\b', "Funding Stage"),
        (r'founded\s+in\s+\d{4}', "Founded"),
        (r'(\d[\d,]+)\s+employees', "Team Size"),
        (r'(\d[\d,]+)\+?\s+(users|customers|clients)', "User Base"),
    ]

    metrics: list[tuple[str, str]] = []
    all_content = " ".join(
        s.get("content", "") for s in sections.values() if isinstance(s, dict)
    )
    for pattern, label in metric_patterns:
        m = re.search(pattern, all_content, re.IGNORECASE)
        if m and len(metrics) < 5:
            metrics.append((label, m.group(0).strip()))

    if metrics:
        lines.append("| Metric | Value |")
        lines.append("|---|---|")
        for label, value in metrics:
            lines.append(f"| {label} | {value} |")
        lines.append("")

    # ── 3. Notable Highlights ───────────────────────────────────────────────
    priority_keys = (
        ["financials", "news", "products", "leadership"]
        if report_type == "company"
        else ["financials", "trends", "news", "analyst_sentiment"]
    )
    all_highlights: list[str] = []
    for key in priority_keys:
        if key in sections and isinstance(sections[key], dict):
            highlights = _extract_key_highlights(sections[key].get("content", ""), max_points=2)
            all_highlights.extend(highlights)

    if all_highlights:
        lines.append("**Key Highlights**")
        lines.append("")
        for h in all_highlights[:6]:
            lines.append(f"- {h}")
        lines.append("")

    # ── 4. Recent Developments ──────────────────────────────────────────────
    for key in ("news", "analyst_sentiment"):
        if key in sections and isinstance(sections[key], dict):
            news_content = sections[key].get("content", "")
            for para in news_content.split("\n"):
                para = para.strip()
                if len(para) > 80 and not para.startswith('#') and not para.startswith('|'):
                    lines.append(f"> {para[:400]}")
                    break
            break

    return "\n".join(lines)


def _render_section_block_v2(key: str, display_name: str, section: dict) -> list[str]:
    """Render a research section with: heading, key highlights callout, full content, sources."""
    emoji = _section_emoji(key)
    lines: list[str] = []

    lines.append(f"## {emoji} {display_name}")
    lines.append("")

    content = _clean_content_for_doc(section.get("content", ""))

    # Key highlights callout — pull top 3 signal sentences
    highlights = _extract_key_highlights(content, max_points=3)
    if highlights:
        lines.append("> **Key Highlights**")
        for h in highlights:
            lines.append(f">")
            lines.append(f"> - {h}")
        lines.append("")

    if content:
        # Demote nested headings so they sit below the ## section header
        content = re.sub(
            r'^(#{1,4})\s',
            lambda m: '#' * min(len(m.group(1)) + 2, 6) + ' ',
            content, flags=re.MULTILINE,
        )
        lines.append(content)
    else:
        lines.append("*No data available for this section.*")

    lines.append("")

    sources = _clean_sources(section.get("sources", []))
    if sources:
        lines.append("<details>")
        lines.append(f"<summary>📎 Sources ({len(sources)})</summary>")
        lines.append("")
        for i, s in enumerate(sources, 1):
            if s["url"]:
                lines.append(f"{i}. [{s['title']}]({s['url']})")
            else:
                lines.append(f"{i}. {s['title']}")
        lines.append("")
        lines.append("</details>")
        lines.append("")

    lines.append("---")
    lines.append("")
    return lines


def create_company_research_doc(company_data: dict, title: str = "") -> str:
    """Create a professional company research report from raw research data.

    Input: dict from research_company() or similar
    Output: Formatted markdown document
    """
    _ensure_dirs()

    company = company_data.get("company", "Unknown Company")
    doc_title = title or f"{company} — Company Research Report"
    sections = company_data.get("sections", {})
    timestamp = datetime.now()

    SECTION_ORDER = [
        ("overview",    "Company Overview"),
        ("leadership",  "Leadership & Team"),
        ("financials",  "Funding & Financials"),
        ("products",    "Products & Technology"),
        ("news",        "Recent News & Announcements"),
        ("competitors", "Competitive Landscape"),
    ]
    present_sections = [(k, n) for k, n in SECTION_ORDER if k in sections]

    lines: list[str] = []

    # ── Cover Page ───────────────────────────────────────────────────────────
    lines.append(f"# {doc_title}")
    lines.append("")
    lines.append("| | |")
    lines.append("|---|---|")
    lines.append(f"| **Company** | {company} |")
    lines.append(f"| **Report Type** | Company Deep Research |")
    lines.append(f"| **Date** | {timestamp.strftime('%B %d, %Y')} |")
    lines.append(f"| **Prepared By** | Clawspan Research Engine |")
    lines.append(f"| **Sections** | {len(present_sections)} |")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ── Table of Contents ────────────────────────────────────────────────────
    lines.append("## Contents")
    lines.append("")
    lines.append("1. [Executive Summary](#executive-summary)")
    for i, (key, display_name) in enumerate(present_sections, 2):
        emoji = _section_emoji(key)
        anchor = display_name.lower().replace(' ', '-').replace('&', '').replace('--', '-')
        lines.append(f"{i}. [{emoji} {display_name}](#{anchor})")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ── Executive Summary ────────────────────────────────────────────────────
    lines.append("## Executive Summary")
    lines.append("")
    smart_summary = _build_smart_exec_summary(company, sections, report_type="company")
    if smart_summary.strip():
        lines.append(smart_summary)
    else:
        lines.append(f"**{company}** is the subject of this research report. "
                     f"The report covers the company's overview, leadership, financials, "
                     f"products and technology, recent news, and competitive landscape.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ── Research Sections ────────────────────────────────────────────────────
    for key, display_name in present_sections:
        lines.extend(_render_section_block_v2(key, display_name, sections[key]))

    # ── Footer ───────────────────────────────────────────────────────────────
    lines.append(f"*Compiled by Clawspan · {timestamp.strftime('%B %d, %Y at %I:%M %p')} · Powered by Tavily*")
    lines.append("")

    return "\n".join(lines)


def create_market_analysis_doc(market_data: dict, title: str = "") -> str:
    """Create a market analysis report from raw research data."""
    _ensure_dirs()

    subject = market_data.get("subject", "Unknown")
    doc_title = title or f"{subject} — Market Analysis"
    sections = market_data.get("sections", {})
    timestamp = datetime.now()

    SECTION_ORDER = [
        ("overview",          "Market Overview"),
        ("financials",        "Financial Performance"),
        ("trends",            "Market Trends & Outlook"),
        ("competitors",       "Competitor Analysis"),
        ("analyst_sentiment", "Analyst Sentiment"),
        ("news",              "Recent Developments"),
    ]
    present_sections = [(k, n) for k, n in SECTION_ORDER if k in sections]

    lines: list[str] = []

    # ── Cover Page ───────────────────────────────────────────────────────────
    lines.append(f"# {doc_title}")
    lines.append("")
    lines.append("| | |")
    lines.append("|---|---|")
    lines.append(f"| **Subject** | {subject} |")
    lines.append(f"| **Report Type** | Market Analysis |")
    lines.append(f"| **Date** | {timestamp.strftime('%B %d, %Y')} |")
    lines.append(f"| **Prepared By** | Clawspan Research Engine |")
    lines.append(f"| **Sections** | {len(present_sections)} |")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ── Table of Contents ────────────────────────────────────────────────────
    lines.append("## Contents")
    lines.append("")
    lines.append("1. [Executive Summary](#executive-summary)")
    for i, (key, display_name) in enumerate(present_sections, 2):
        emoji = _section_emoji(key)
        anchor = display_name.lower().replace(' ', '-').replace('&', '').replace('--', '-')
        lines.append(f"{i}. [{emoji} {display_name}](#{anchor})")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ── Executive Summary ────────────────────────────────────────────────────
    lines.append("## Executive Summary")
    lines.append("")
    smart_summary = _build_smart_exec_summary(subject, sections, report_type="market")
    if smart_summary.strip():
        lines.append(smart_summary)
    else:
        lines.append(f"This report provides a comprehensive market analysis of **{subject}**, "
                     f"covering market overview, financial performance, trends, competitive "
                     f"dynamics, and recent developments.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ── Research Sections ────────────────────────────────────────────────────
    for key, display_name in present_sections:
        lines.extend(_render_section_block_v2(key, display_name, sections[key]))

    # ── Footer ───────────────────────────────────────────────────────────────
    lines.append(f"*Compiled by Clawspan · {timestamp.strftime('%B %d, %Y at %I:%M %p')} · Powered by Tavily*")
    lines.append("")

    return "\n".join(lines)


def create_meeting_prep_doc(meeting_data: dict, title: str = "") -> str:
    """Create a meeting preparation document."""
    _ensure_dirs()

    sections = meeting_data.get("sections", {})
    company = ""
    attendee = ""
    topic = ""
    for sec_name, sec_data in sections.items():
        if "company" in sec_name.lower():
            company = sec_data.get("name", "").replace("Company: ", "")
        elif "attendee" in sec_name.lower():
            attendee = sec_data.get("name", "").replace("Attendee: ", "")
        elif "topic" in sec_name.lower():
            topic = sec_data.get("name", "").replace("Topic: ", "")

    doc_title = title or f"Meeting Prep: {company or 'Unknown'}"

    lines = []
    lines.append(f"# {doc_title}")
    lines.append("")
    lines.append(f"> **Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    if company: lines.append(f"> **Company:** {company}")
    if attendee: lines.append(f"> **Attendee:** {attendee}")
    if topic: lines.append(f"> **Topic:** {topic}")
    lines.append(f"> **Type:** Meeting Preparation")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Table of Contents
    lines.append("## Table of Contents")
    lines.append("")
    for sec_data in sections.values():
        name = sec_data.get("name", "")
        anchor = name.lower().replace(" ", "-")
        lines.append(f"- [{name}](#{anchor})")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Sections
    for sec_data in sections.values():
        name = sec_data.get("name", "")
        content = sec_data.get("content", "")
        lines.append(f"## {name}")
        lines.append("")
        content_clean = re.sub(r'^#+\s*.*$', '', content, flags=re.MULTILINE).strip()
        lines.append(content_clean[:2000])
        lines.append("")
        lines.append("---")
        lines.append("")

    lines.append("")
    lines.append(f"*Prepared by Clawspan Writer on {datetime.now().strftime('%B %d, %Y')}*")
    lines.append("")

    return "\n".join(lines)


def create_technical_doc(content: str, title: str = "", doc_type: str = "README") -> str:
    """Create technical documentation (README, API docs, architecture docs)."""
    _ensure_dirs()

    doc_title = title or f"{doc_type} Document"

    if doc_type.upper() == "README":
        lines = []
        lines.append(f"# {doc_title}")
        lines.append("")
        lines.append(content)
        lines.append("")
        lines.append("---")
        lines.append(f"*Generated by Clawspan Writer on {datetime.now().strftime('%B %d, %Y')}*")
        lines.append("")
        return "\n".join(lines)

    # Generic technical doc
    lines = []
    lines.append(f"# {doc_title}")
    lines.append("")
    lines.append(f"> **Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    lines.append(f"> **Type:** {doc_type}")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(content)
    lines.append("")
    return "\n".join(lines)


def create_custom_doc(title: str, content: str, doc_type: str = "document") -> str:
    """Create a custom document with the given content."""
    _ensure_dirs()

    lines = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"> **Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    lines.append(f"> **Type:** {doc_type.title()}")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(content)
    lines.append("")
    lines.append("---")
    lines.append(f"*Generated by Clawspan Writer on {datetime.now().strftime('%B %d, %Y')}*")
    lines.append("")
    return "\n".join(lines)


# ── Save Document ─────────────────────────────────────────────────────────────

def save_document(content: str, title: str, doc_type: str = "document",
                  folder: str = "") -> str:
    """Save a markdown document to the appropriate folder.

    Returns the full file path.
    """
    _ensure_dirs()

    subfolder = folder or _get_subfolder(doc_type)
    filename = _generate_filename(title, doc_type)
    filepath = os.path.join(DOCS_DIR, subfolder, filename)

    with open(filepath, "w") as f:
        f.write(content)

    return filepath


# ── Export Document ──────────────────────────────────────────────────────────

def export_document(source_path: str, output_format: str = "pdf") -> str:
    """Export a markdown document to another format.

    Formats: pdf, docx, html, txt

    Returns the exported file path.
    """
    if not os.path.exists(source_path):
        return f"Source file not found: {source_path}"

    base = os.path.splitext(source_path)[0]

    if output_format == "pdf":
        output_path = f"{base}.pdf"
        return _export_to_pdf(source_path, output_path)
    elif output_format == "docx":
        output_path = f"{base}.docx"
        return _export_to_docx(source_path, output_path)
    elif output_format == "html":
        output_path = f"{base}.html"
        return _export_to_html(source_path, output_path)
    elif output_format == "txt":
        output_path = f"{base}.txt"
        return _export_to_txt(source_path, output_path)

    return f"Unsupported format: {output_format}. Use: pdf, docx, html, txt"


def _export_to_pdf(source: str, output: str) -> str:
    """Export markdown to PDF via weasyprint."""
    try:
        # Convert md → HTML first (using pandoc for clean output)
        html_content = _md_to_html(source)
        styled_html = _add_pdf_styles(html_content)

        from weasyprint import HTML
        HTML(string=styled_html).write_pdf(output)
        return f"PDF saved: {output}"
    except Exception as e:
        return f"PDF export error: {e}"


def _add_inline_runs(paragraph, text: str) -> None:
    """Parse inline markdown (bold, italic, links) and add styled runs to a paragraph."""
    # Strip markdown links → just the display text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Split on **bold** and *italic* markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_hr(doc) -> None:
    """Add a thin horizontal rule paragraph."""
    import docx as _docx
    from docx.oxml.ns import qn
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = _docx.oxml.OxmlElement('w:pBdr')
    bottom = _docx.oxml.OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_table(doc, table_lines: list[str]) -> None:
    """Render a markdown pipe table into a proper DOCX table."""
    import docx as _docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    # Parse rows — skip separator lines (---|---)
    rows = []
    for line in table_lines:
        line = line.strip().strip('|')
        if re.match(r'^[\s\-|:]+$', line):
            continue
        cells = [c.strip() for c in line.split('|')]
        if cells:
            rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    # Pad short rows
    rows = [r + [''] * (max_cols - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=max_cols)
    tbl.style = 'Table Grid'

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            # Strip remaining markdown from cell text
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
            clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
            p = cell.paragraphs[0]
            run = p.add_run(clean)
            run.font.size = Pt(10)
            if r_idx == 0:
                run.bold = True
                # Header row shading
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = _docx.oxml.OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1F3864')
                tcPr.append(shd)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_cover_title(doc, title: str) -> None:
    """Render the document H1 as a full-width dark blue cover title block."""
    import docx as _docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)

    # Dark navy background shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = _docx.oxml.OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F3864')
    pPr.append(shd)

    run = p.add_run(title)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _render_cover_table(doc, table_lines: list[str]) -> None:
    """Render the cover metadata table — left col bold dark, right col value, no header row."""
    import docx as _docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    rows = []
    for line in table_lines:
        line = line.strip().strip('|')
        if re.match(r'^[\s\-|:]+$', line):
            continue
        cells = [c.strip() for c in line.split('|')]
        if len(cells) >= 2 and any(c for c in cells):
            rows.append(cells[:2])

    if not rows:
        return

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'

    for r_idx, (label, value) in enumerate(rows):
        # Left cell — label
        label_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', label).strip()
        lc = tbl.cell(r_idx, 0)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label_clean)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        # Shading on label cell
        tcPr = lc._tc.get_or_add_tcPr()
        shd = _docx.oxml.OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EBF0FA')
        tcPr.append(shd)

        # Right cell — value
        value_clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', value).strip()
        vc = tbl.cell(r_idx, 1)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value_clean)
        vr.font.size = Pt(10)


def _render_callout_block(doc, lines: list[str]) -> None:
    """Render a > blockquote block as a styled left-bordered callout box."""
    import docx as _docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    for line in lines:
        inner = line.lstrip('> ').strip()
        if not inner:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(2)

        # Left border to simulate callout
        pPr = p._p.get_or_add_pPr()
        pBdr = _docx.oxml.OxmlElement('w:pBdr')
        left = _docx.oxml.OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '2E74B5')
        pBdr.append(left)
        pPr.append(pBdr)

        # Shading
        shd = _docx.oxml.OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EFF4FB')
        pPr.append(shd)

        _add_inline_runs(p, inner)
        for run in p.runs:
            run.font.size = Pt(10)
            if not run.bold:
                run.font.color.rgb = RGBColor(0x24, 0x49, 0x7A)


def _export_to_docx(source: str, output: str) -> str:
    """Export markdown to a clean, properly formatted DOCX — no raw markdown symbols."""
    try:
        import docx as _docx
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = _docx.Document()

        # ── Page margins ─────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)

        # ── Document-wide typography ─────────────────────────────────────
        for style_name in ('Normal', 'Body Text'):
            try:
                s = doc.styles[style_name]
                s.font.name = 'Calibri'
                s.font.size = Pt(11)
            except Exception:
                pass

        # Heading styles
        heading_colors = {1: '1F3864', 2: '1F3864', 3: '2E74B5', 4: '404040'}
        heading_sizes  = {1: 22, 2: 16, 3: 13, 4: 12}
        for level, color_hex in heading_colors.items():
            try:
                h = doc.styles[f'Heading {level}']
                h.font.name = 'Calibri'
                h.font.size = Pt(heading_sizes[level])
                h.font.color.rgb = RGBColor(
                    int(color_hex[0:2], 16),
                    int(color_hex[2:4], 16),
                    int(color_hex[4:6], 16),
                )
            except Exception:
                pass

        with open(source, encoding='utf-8') as f:
            content = f.read()

        lines = content.split('\n')
        i = 0
        first_h1_done = False  # track whether we've rendered the cover title

        while i < len(lines):
            line = lines[i].rstrip()

            # ── Skip TOC anchor links ─────────────────────────────────────
            if re.match(r'^\d+\.\s+\[.*\]\(#.*\)\s*$', line):
                i += 1
                continue
            if re.match(r'^-\s+\[.*\]\(#.*\)\s*$', line):
                i += 1
                continue

            # ── <details> / sources block ─────────────────────────────────
            if re.match(r'^\s*<details>', line, re.IGNORECASE):
                i += 1
                sources_lines = []
                while i < len(lines) and not re.match(r'^\s*</details>', lines[i], re.IGNORECASE):
                    l = lines[i].strip()
                    if re.match(r'<summary>.*</summary>', l, re.IGNORECASE):
                        i += 1
                        continue
                    m = re.match(r'^(\d+)\.\s+\[([^\]]+)\]\(([^)]+)\)', l)
                    if m:
                        sources_lines.append(f"{m.group(1)}. {m.group(2)}")
                    i += 1
                if sources_lines:
                    p = doc.add_paragraph()
                    r = p.add_run('Sources:')
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                    for sl in sources_lines:
                        sp = doc.add_paragraph(sl, style='Normal')
                        if sp.runs:
                            sp.runs[0].font.size = Pt(9)
                            sp.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                i += 1
                continue

            # ── Skip bare HTML tags ───────────────────────────────────────
            if re.match(r'^\s*<[^>]+>\s*$', line):
                i += 1
                continue

            # ── Empty line ────────────────────────────────────────────────
            if not line.strip():
                doc.add_paragraph('')
                i += 1
                continue

            # ── Horizontal rule ───────────────────────────────────────────
            if re.match(r'^---+$', line.strip()) or re.match(r'^\*\*\*+$', line.strip()):
                _add_hr(doc)
                i += 1
                continue

            # ── Cover metadata table (2 columns, no header row shading) ───
            # Detect: first table that appears after the H1, before any H2
            if '|' in line and re.match(r'^\s*\|', line) and not first_h1_done:
                # still in cover area — skip to normal table render
                pass  # fall through to table handler below

            # ── Markdown table ─────────────────────────────────────────────
            if '|' in line and re.match(r'^\s*\|', line):
                table_lines = []
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                # Cover metadata table: 2 cols, appears right after H1, no header highlight
                is_cover_table = not first_h1_done and len(table_lines) >= 2
                if is_cover_table:
                    _render_cover_table(doc, table_lines)
                else:
                    _render_table(doc, table_lines)
                continue

            # ── H1 → styled cover title ───────────────────────────────────
            if re.match(r'^#\s+', line) and not first_h1_done:
                title_text = re.sub(r'^#\s+', '', line)
                title_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', title_text).strip()
                _add_cover_title(doc, title_text)
                first_h1_done = True
                i += 1
                continue

            # ── Other headings ────────────────────────────────────────────
            header_match = re.match(r'^(#{1,6})\s+(.*)', line)
            if header_match:
                level = min(len(header_match.group(1)), 4)
                text = header_match.group(2)
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                text = re.sub(r'\*([^*]+)\*', r'\1', text)
                text = text.strip()
                if text:
                    doc.add_heading(text, level=level)
                i += 1
                continue

            # ── Blockquote callout block ──────────────────────────────────
            # Collect all consecutive > lines and render as callout
            if line.startswith('> ') or line == '>':
                callout_lines = []
                while i < len(lines) and (lines[i].startswith('> ') or lines[i].strip() == '>'):
                    callout_lines.append(lines[i])
                    i += 1
                _render_callout_block(doc, callout_lines)
                continue

            # ── Bullet list ───────────────────────────────────────────────
            if re.match(r'^[-*]\s+', line):
                text = re.sub(r'^[-*]\s+', '', line)
                p = doc.add_paragraph(style='List Bullet')
                _add_inline_runs(p, text)
                p.paragraph_format.space_after = Pt(2)
                i += 1
                continue

            # ── Numbered list ─────────────────────────────────────────────
            if re.match(r'^\d+\.\s+', line):
                text = re.sub(r'^\d+\.\s+', '', line)
                p = doc.add_paragraph(style='List Number')
                _add_inline_runs(p, text)
                p.paragraph_format.space_after = Pt(2)
                i += 1
                continue

            # ── Footer italic line ────────────────────────────────────────
            if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
                p = doc.add_paragraph(line.strip('*'))
                if p.runs:
                    p.runs[0].italic = True
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                i += 1
                continue

            # ── Normal paragraph ──────────────────────────────────────────
            p = doc.add_paragraph()
            _add_inline_runs(p, line)
            p.paragraph_format.space_after = Pt(4)
            i += 1

        doc.save(output)
        return f"DOCX saved: {output}"
    except Exception as e:
        return f"DOCX export error: {e}"


def _export_to_html(source: str, output: str) -> str:
    """Export markdown to HTML."""
    try:
        html = _md_to_html(source)
        styled = _add_web_styles(html)
        with open(output, "w") as f:
            f.write(styled)
        return f"HTML saved: {output}"
    except Exception as e:
        return f"HTML export error: {e}"


def _export_to_txt(source: str, output: str) -> str:
    """Export markdown to plain text (strip formatting)."""
    try:
        with open(source) as f:
            content = f.read()
        # Strip markdown formatting
        text = re.sub(r'#{1,6}\s+', '', content)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        text = re.sub(r'^>\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
        text = re.sub(r'\n{3,}', '\n\n', text).strip()
        with open(output, "w") as f:
            f.write(text)
        return f"Text saved: {output}"
    except Exception as e:
        return f"Text export error: {e}"


def _md_to_html(source: str) -> str:
    """Convert markdown to HTML via pandoc (or simple regex fallback)."""
    try:
        result = subprocess.run(
            ["pandoc", source, "-f", "markdown", "-t", "html"],
            capture_output=True, text=True, timeout=15,
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Simple regex-based conversion
    with open(source) as f:
        content = f.read()

    lines = content.split("\n")
    html_lines = []

    for line in lines:
        # Headers
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            html_lines.append(f"<h{level}>{m.group(2)}</h{level}>")
            continue

        # Blockquotes
        if line.startswith("> "):
            html_lines.append(f"<blockquote>{line[2:]}</blockquote>")
            continue

        # Bold/italic
        line = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', line)
        line = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', line)

        # Links
        line = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', line)

        # Horizontal rules
        if line.strip() in ("---", "***"):
            html_lines.append("<hr>")
            continue

        # Paragraphs
        if line.strip():
            html_lines.append(f"<p>{line}</p>")
        else:
            html_lines.append("")

    return "\n".join(html_lines)


def _add_pdf_styles(html: str) -> str:
    """Add clean print-ready styles for PDF export."""
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{
    font-family: 'Georgia', serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
    max-width: 48rem;
    margin: 0 auto;
    padding: 2rem;
}}
h1 {{ font-size: 24pt; color: #1a1a2e; border-bottom: 2px solid #16213e; padding-bottom: 0.3rem; }}
h2 {{ font-size: 18pt; color: #16213e; border-bottom: 1px solid #ddd; padding-bottom: 0.2rem; margin-top: 2rem; }}
h3 {{ font-size: 14pt; color: #0f3460; }}
blockquote {{
    border-left: 3px solid #16213e;
    margin-left: 0;
    padding-left: 1rem;
    color: #666;
    font-style: italic;
}}
a {{ color: #0f3460; }}
hr {{ border: none; border-top: 1px solid #ddd; margin: 2rem 0; }}
</style>
</head>
<body>
{html}
</body>
</html>"""


def _add_web_styles(html: str) -> str:
    """Add clean web styles for HTML export."""
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    font-size: 16px;
    line-height: 1.7;
    color: #24292e;
    max-width: 800px;
    margin: 40px auto;
    padding: 0 20px;
}}
h1 {{ font-size: 2em; border-bottom: 2px solid #eaecef; padding-bottom: 0.3em; }}
h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
h3 {{ font-size: 1.25em; }}
blockquote {{
    border-left: 4px solid #dfe2e5;
    margin: 0;
    padding: 0 1em;
    color: #6a737d;
}}
code {{
    background: #f6f8fa;
    padding: 0.2em 0.4em;
    border-radius: 3px;
    font-size: 85%;
}}
hr {{ border: 0; border-top: 1px solid #eaecef; margin: 2em 0; }}
a {{ color: #0366d6; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
</style>
</head>
<body>
{html}
</body>
</html>"""


# ── Document Editing ─────────────────────────────────────────────────────────

def edit_document(file_path: str, edit_type: str, value: str = "") -> str:
    """Edit an existing document.

    edit_type: "shorten", "expand", "add_toc", "to_bullets", "to_email",
               "append", "prepend", "replace_section"
    """
    if not os.path.exists(file_path):
        return f"File not found: {file_path}"

    with open(file_path) as f:
        content = f.read()

    if edit_type == "shorten":
        result = _shorten_document(content)
    elif edit_type == "add_toc":
        result = _add_toc_to_document(content)
    elif edit_type == "to_bullets":
        result = _convert_to_bullets(content)
    elif edit_type == "to_email":
        result = _convert_to_email(content, value)
    elif edit_type == "append":
        result = content + "\n\n" + value
    elif edit_type == "prepend":
        result = value + "\n\n" + content
    elif edit_type == "replace_section":
        # value should be "OLD_TEXT|||NEW_TEXT"
        if "|||" in value:
            old, new = value.split("|||", 1)
            result = content.replace(old.strip(), new.strip())
        else:
            return "For replace_section, value should be: old_text|||new_text"
    else:
        return f"Unknown edit type: {edit_type}"

    # Overwrite the file
    with open(file_path, "w") as f:
        f.write(result)

    return f"Document updated: {file_path}"


def _shorten_document(content: str) -> str:
    """Condense document to key points — keeps headers and first paragraph of each section."""
    lines = content.split("\n")
    result = []
    in_paragraph = False
    paragraph_lines = 0

    for line in lines:
        # Keep headers
        if re.match(r'^#{1,6}\s+', line):
            result.append(line)
            in_paragraph = False
            paragraph_lines = 0
            continue

        # Keep first ~3 lines after each header
        if not in_paragraph and line.strip():
            result.append(line)
            paragraph_lines += 1
            in_paragraph = True
            if paragraph_lines >= 3:
                result.append("*...[condensed]*")
            continue

        if line.strip() == "" or line.strip() == "---":
            result.append(line)
            in_paragraph = False
            paragraph_lines = 0
            continue

        if in_paragraph and paragraph_lines < 3:
            result.append(line)
            paragraph_lines += 1

    return "\n".join(result)


def _add_toc_to_document(content: str) -> str:
    """Add a table of contents to a document (after the first header)."""
    lines = content.split("\n")
    toc_lines = []
    first_header_idx = None

    for i, line in enumerate(lines):
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            if first_header_idx is None:
                first_header_idx = i
            level = len(m.group(1))
            if level <= 3:  # Only H1-H3 in TOC
                title = m.group(2)
                anchor = title.lower()
                anchor = re.sub(r'[^\w\s-]', '', anchor)
                anchor = anchor.replace(" ", "-")
                indent = "  " * (level - 1)
                toc_lines.append(f"{indent}- [{title}](#{anchor})")

    if first_header_idx is not None:
        # Find where to insert TOC (after first blockquote/metadata section)
        insert_idx = first_header_idx + 1
        while insert_idx < len(lines) and (lines[insert_idx].startswith(">") or lines[insert_idx].strip() == "---" or not lines[insert_idx].strip()):
            insert_idx += 1

        toc_content = "\n\n## Table of Contents\n\n" + "\n".join(toc_lines) + "\n"
        lines.insert(insert_idx, toc_content)

    return "\n".join(lines)


def _convert_to_bullets(content: str) -> str:
    """Convert paragraph-heavy sections to bullet points."""
    lines = content.split("\n")
    result = []

    for line in lines:
        # Keep headers, blockquotes, horizontal rules
        if re.match(r'^#{1,6}\s+', line) or line.startswith(">") or line.strip() == "---" or not line.strip():
            result.append(line)
            continue

        # Convert non-empty, non-header lines to bullets
        if line.strip() and not line.startswith("-") and not line.startswith("*"):
            # Only if it looks like a paragraph (not already a list or code)
            if not line.startswith(" ") and not line.startswith("```"):
                result.append(f"- {line}")
            else:
                result.append(line)
        else:
            result.append(line)

    return "\n".join(result)


def _convert_to_email(content: str, to_address: str = "") -> str:
    """Convert document content to email format."""
    # Extract title (first header)
    title = "Document"
    m = re.match(r'^#\s+(.*)', content)
    if m:
        title = m.group(1)

    # Strip the header and metadata
    body = re.sub(r'^#\s+.*?\n', '', content, count=1)
    body = re.sub(r'^>.*\n', '', body, flags=re.MULTILINE)
    body = re.sub(r'^---\s*\n', '', body)
    body = body.strip()

    lines = [
        f"Subject: {title}",
    ]
    if to_address:
        lines.append(f"To: {to_address}")
    lines.append("")
    lines.append(body)
    lines.append("")

    return "\n".join(lines)


# ── List & Search Documents ──────────────────────────────────────────────────

def list_documents(doc_type: str = "", limit: int = 20) -> str:
    """List saved documents, optionally filtered by type."""
    _ensure_dirs()

    if doc_type:
        folders = [_get_subfolder(doc_type)]
    else:
        folders = list(DOC_SUBFOLDERS.values())

    all_files = []
    for folder in folders:
        folder_path = os.path.join(DOCS_DIR, folder)
        if os.path.exists(folder_path):
            for f in sorted(os.listdir(folder_path), reverse=True):
                if f.endswith((".md", ".pdf", ".docx", ".html", ".txt")):
                    all_files.append((folder, f, os.path.join(folder_path, f)))

    if not all_files:
        return "No documents saved yet."

    all_files.sort(key=lambda x: x[2], reverse=True)
    all_files = all_files[:limit]

    lines = [f"Saved Documents ({len(all_files)}):"]
    lines.append("")
    for folder, filename, filepath in all_files:
        size = os.path.getsize(filepath)
        size_str = f"{size/1024:.1f}KB" if size > 1024 else f"{size}B"
        lines.append(f"📄 [{folder}] {filename} ({size_str})")
        lines.append(f"   {filepath}")

    return "\n".join(lines)


def get_document_info(file_path: str) -> str:
    """Get information about a saved document."""
    if not os.path.exists(file_path):
        return f"File not found: {file_path}"

    stat = os.stat(file_path)
    size = stat.st_size
    created = datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M")
    modified = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M")

    ext = os.path.splitext(file_path)[1]
    with open(file_path, encoding="utf-8", errors="replace") as f:
        content = f.read()

    word_count = len(content.split())
    line_count = len(content.split("\n"))

    # Extract title (first header)
    title = "Unknown"
    m = re.match(r'^#\s+(.*)', content)
    if m:
        title = m.group(1)

    return (
        f"Document: {title}\n"
        f"File: {file_path}\n"
        f"Type: {ext}\n"
        f"Size: {size/1024:.1f} KB\n"
        f"Words: {word_count}\n"
        f"Lines: {line_count}\n"
        f"Created: {created}\n"
        f"Modified: {modified}"
    )


def delete_document(file_path: str) -> str:
    """Delete a saved document."""
    if not os.path.exists(file_path):
        return f"File not found: {file_path}"

    os.remove(file_path)
    return f"Deleted: {file_path}"


def read_document(file_path: str, max_chars: int = 5000) -> str:
    """Read a saved document's contents."""
    if not os.path.exists(file_path):
        return f"File not found: {file_path}"

    with open(file_path, encoding="utf-8", errors="replace") as f:
        content = f.read()

    if len(content) > max_chars:
        return content[:max_chars] + f"\n\n...[truncated, {len(content) - max_chars} more chars]"
    return content
